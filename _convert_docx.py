import traceback
try:
    from docx import Document
    d = Document(r'D:\Onedrive\桌面\杨志诚毕业论文1.docx')
    with open('thesis_full.txt', 'w', encoding='utf-8') as f:
        for i, p in enumerate(d.paragraphs):
            f.write(f'{i:04d}\t{p.text}\n\n')
    print(f'Done, paragraphs: {len(d.paragraphs)}')
except Exception as e:
    traceback.print_exc()
