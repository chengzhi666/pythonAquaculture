"""Apply the 2025 CAU thesis formatting rules to a Word thesis document.

The script is intentionally conservative:
- it keeps the document structure and content intact;
- it normalizes page setup, styles, headers/footers, and common paragraph types;
- it repairs broken `NULL` relationships that can stop `python-docx` from opening
  the file.

Example:
    python fix_thesis_format.py --input thesis.docx --output thesis_cau_2025.docx
"""

from __future__ import annotations

import argparse
import copy
import os
import tempfile
import zipfile
from pathlib import Path

from lxml import etree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


SPECIAL_CENTER_HEADINGS = {
    "摘  要",
    "Abstract",
    "目   录",
    "目 录",
    "插图清单",
    "附表清单",
    "主要符号表",
}
DECLARATION_HEADINGS = {
    "独 创 性 声 明",
    "关于学位论文使用授权的说明",
}
KEYWORD_PREFIXES = ("关键词：", "Key words:")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Path to the source .docx")
    parser.add_argument("--output", required=True, help="Path to the formatted .docx")
    parser.add_argument(
        "--header-left",
        default=None,
        help="Override the main-text header left text, for example 中国农业大学硕士学位论文",
    )
    return parser.parse_args()


def set_rfonts(target, cn_font: str, en_font: str) -> None:
    rfonts = target.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        target.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:cs"), en_font)


def set_run_fonts(run, cn_font: str, en_font: str, size=None, bold=None) -> None:
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = en_font
    set_rfonts(run._element.get_or_add_rPr(), cn_font, en_font)


def set_style_fonts(style, cn_font: str, en_font: str, size=None, bold=None) -> None:
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    style.font.name = en_font
    set_rfonts(style.element.get_or_add_rPr(), cn_font, en_font)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def reset_story(story):
    element = story._element
    for child in list(element):
        element.remove(child)
    paragraph = OxmlElement("w:p")
    element.append(paragraph)
    return story.paragraphs[0]


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_field(paragraph, instruction: str, font_size=Pt(9), placeholder: str = "1") -> None:
    begin_run = paragraph.add_run()
    set_run_fonts(begin_run, "宋体", "Times New Roman", font_size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    set_run_fonts(instr_run, "宋体", "Times New Roman", font_size)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    set_run_fonts(sep_run, "宋体", "Times New Roman", font_size)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    text_run = paragraph.add_run(placeholder)
    set_run_fonts(text_run, "宋体", "Times New Roman", font_size)

    end_run = paragraph.add_run()
    set_run_fonts(end_run, "宋体", "Times New Roman", font_size)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_bottom_border(paragraph, line_type: str = "double", size_eighth_pt: int = 24) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    border = ppr.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        ppr.append(border)
    else:
        for child in list(border):
            border.remove(child)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), line_type)
    bottom.set(qn("w:sz"), str(size_eighth_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)


def set_page_number_type(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is None:
        if qn("w:start") in pg_num_type.attrib:
            del pg_num_type.attrib[qn("w:start")]
    else:
        pg_num_type.set(qn("w:start"), str(start))


def normalize_run_family(run, cn_font: str, en_font: str) -> None:
    run.font.name = en_font
    set_rfonts(run._element.get_or_add_rPr(), cn_font, en_font)


def normalize_run_family_keep_size(run, cn_font: str, en_font: str, bold=None) -> None:
    size = run.font.size
    set_run_fonts(run, cn_font, en_font, size=size, bold=run.font.bold if bold is None else bold)


def format_runs_with_size_guard(
    paragraph,
    cn_font: str,
    en_font: str,
    default_size,
    *,
    default_bold=None,
    min_keep_pt: float | None = None,
    max_keep_pt: float | None = None,
) -> None:
    for run in paragraph.runs:
        current = run.font.size.pt if run.font.size else None
        if current is not None:
            if min_keep_pt is not None and current < min_keep_pt:
                normalize_run_family_keep_size(run, cn_font, en_font)
                continue
            if max_keep_pt is not None and current > max_keep_pt:
                normalize_run_family_keep_size(run, cn_font, en_font)
                continue
        set_run_fonts(
            run,
            cn_font,
            en_font,
            size=current and Pt(current) or default_size,
            bold=run.font.bold if default_bold is None else default_bold,
        )


def apply_heading_style(paragraph, size_pt: float, align, *, bold: bool = True) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(size_pt)
    paragraph.paragraph_format.space_after = Pt(size_pt)
    paragraph.paragraph_format.line_spacing = None
    format_runs_with_size_guard(
        paragraph,
        "黑体",
        "Times New Roman",
        Pt(size_pt),
        default_bold=bold,
        min_keep_pt=10.0,
        max_keep_pt=26.0,
    )
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def apply_center_heading(paragraph, size_pt: float) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    for run in paragraph.runs:
        set_run_fonts(run, "黑体", "Times New Roman", Pt(size_pt), bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)


def apply_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    format_runs_with_size_guard(
        paragraph,
        "宋体",
        "Times New Roman",
        Pt(12),
        min_keep_pt=10.0,
        max_keep_pt=14.0,
    )


def apply_keyword_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for idx, run in enumerate(paragraph.runs):
        set_run_fonts(run, "宋体", "Times New Roman", Pt(12), bold=True if idx == 0 else None)


def apply_list_item_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, "宋体", "Times New Roman", Pt(10.5), bold=run.font.bold)


def apply_cover_title_paragraph(paragraph, *, cn: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    if cn:
        for run in paragraph.runs:
            set_run_fonts(run, "黑体", "Times New Roman", Pt(22), bold=True)
    else:
        for run in paragraph.runs:
            set_run_fonts(run, "宋体", "Times New Roman", Pt(16), bold=True)


def apply_cover_meta_paragraph(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = Pt(20)
    for run in paragraph.runs:
        normalize_run_family_keep_size(run, "宋体", "Times New Roman")


def apply_table_formatting(doc: Document) -> int:
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    for run in paragraph.runs:
                        set_run_fonts(run, "宋体", "Times New Roman", Pt(10.5), bold=run.font.bold)
    return table_count


def fix_docx_null_relationships(src: Path) -> Path:
    tmp = tempfile.NamedTemporaryFile(prefix="cau_fix_", suffix=".docx", delete=False)
    tmp_path = Path(tmp.name)
    tmp.close()

    removed_ids: set[str] = set()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name.endswith(".rels"):
                root = etree.fromstring(data)
                to_remove = []
                for rel in root:
                    target = rel.get("Target", "")
                    if "NULL" in target.upper():
                        removed_ids.add(rel.get("Id"))
                        to_remove.append(rel)
                if to_remove:
                    for rel in to_remove:
                        root.remove(rel)
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)

    if not removed_ids:
        return tmp_path

    with zipfile.ZipFile(tmp_path, "r") as zin, zipfile.ZipFile(f"{tmp_path}.clean", "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name.startswith("word/") and name.endswith(".xml"):
                root = etree.fromstring(data)
                for elem in root.iter():
                    for attr_name, attr_val in list(elem.attrib.items()):
                        if attr_val in removed_ids:
                            del elem.attrib[attr_name]
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(name, data)

    os.replace(f"{tmp_path}.clean", tmp_path)
    return tmp_path


def get_section_ranges(doc: Document) -> list[tuple[int, int]]:
    breaks = []
    for idx, paragraph in enumerate(doc.paragraphs):
        ppr = paragraph._element.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            breaks.append(idx)
    starts = [0] + [idx + 1 for idx in breaks]
    ends = breaks + [len(doc.paragraphs) - 1]
    return list(zip(starts, ends))


def infer_header_left(doc: Document) -> str:
    text = "\n".join(p.text for p in doc.paragraphs[:30])
    if "博士" in text and "硕士" not in text:
        return "中国农业大学博士学位论文"
    return "中国农业大学硕士学位论文"


def looks_like_body_text(text: str) -> bool:
    compact = text.strip()
    if not compact:
        return False
    if compact in SPECIAL_CENTER_HEADINGS or compact in DECLARATION_HEADINGS:
        return False
    if compact.startswith(KEYWORD_PREFIXES):
        return False
    if "\t" in compact:
        return False
    if compact.startswith("第") and "章" in compact[:6]:
        return False
    if compact[:1].isdigit() and "." in compact[:8]:
        return False
    if len(compact) < 24:
        return False
    if "：" in compact[:12] and len(compact) < 40:
        return False
    return True


def section_heading_map(doc: Document, ranges: list[tuple[int, int]]) -> dict[int, str | None]:
    mapping: dict[int, str | None] = {}
    last_title = None
    for sec_idx, (start, end) in enumerate(ranges):
        current = None
        for paragraph in doc.paragraphs[start : end + 1]:
            if paragraph.style and paragraph.style.name == "Heading 1" and paragraph.text.strip():
                current = paragraph.text.strip()
                break
        if current is None:
            current = last_title
        mapping[sec_idx] = current
        if current is not None:
            last_title = current
    return mapping


def configure_headers_and_footers(doc: Document, header_left_text: str) -> tuple[int, int]:
    ranges = get_section_ranges(doc)
    headings = section_heading_map(doc, ranges)

    main_start = None
    for sec_idx, title in headings.items():
        if title and title.startswith("第一章"):
            main_start = sec_idx
            break
    if main_start is None:
        main_start = 0

    front_sections = 0
    main_sections = 0

    doc.settings.odd_and_even_pages_header_footer = False

    for idx, section in enumerate(doc.sections):
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False

        header = reset_story(section.header)
        footer = reset_story(section.footer)
        reset_story(section.first_page_header)
        reset_story(section.first_page_footer)
        reset_story(section.even_page_header)
        reset_story(section.even_page_footer)

        header.paragraph_format.space_before = Pt(0)
        header.paragraph_format.space_after = Pt(0)
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if idx == 0:
            clear_paragraph(header)
            clear_paragraph(footer)
            continue

        if idx < main_start:
            front_sections += 1
            clear_paragraph(header)
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_field(footer, " PAGE ")
            set_page_number_type(section, "upperRoman", start=1 if idx == 1 else None)
            continue

        main_sections += 1
        title = headings.get(idx) or header_left_text
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.first_line_indent = Pt(0)
        header.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )
        left_run = header.add_run(header_left_text)
        set_run_fonts(left_run, "宋体", "Times New Roman", Pt(9))
        header.add_run("\t")
        right_run = header.add_run(title)
        set_run_fonts(right_run, "宋体", "Times New Roman", Pt(9))
        set_bottom_border(header)

        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, " PAGE ")
        set_page_number_type(section, "decimal", start=1 if idx == main_start else None)

    return front_sections, main_sections


def configure_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.header_distance = Cm(2.3)
        section.footer_distance = Cm(1.8)


def configure_styles(doc: Document) -> None:
    if "Normal" in [style.name for style in doc.styles]:
        normal = doc.styles["Normal"]
        set_style_fonts(normal, "宋体", "Times New Roman", Pt(12))
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.line_spacing = Pt(20)
        normal.paragraph_format.first_line_indent = Pt(24)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

    if "Heading 1" in [style.name for style in doc.styles]:
        style = doc.styles["Heading 1"]
        set_style_fonts(style, "黑体", "Times New Roman", Pt(16), bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(16)
        style.paragraph_format.space_after = Pt(16)

    if "Heading 2" in [style.name for style in doc.styles]:
        style = doc.styles["Heading 2"]
        set_style_fonts(style, "黑体", "Times New Roman", Pt(14), bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(14)

    if "Heading 3" in [style.name for style in doc.styles]:
        style = doc.styles["Heading 3"]
        set_style_fonts(style, "黑体", "Times New Roman", Pt(12), bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)

    if "Caption" in [style.name for style in doc.styles]:
        style = doc.styles["Caption"]
        set_style_fonts(style, "黑体", "Times New Roman", Pt(10.5), bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.first_line_indent = Pt(0)

    if "Body Text" in [style.name for style in doc.styles]:
        style = doc.styles["Body Text"]
        set_style_fonts(style, "宋体", "Times New Roman", Pt(10.5))
        style.paragraph_format.first_line_indent = Pt(0)

    if "Body Text Indent" in [style.name for style in doc.styles]:
        style = doc.styles["Body Text Indent"]
        set_style_fonts(style, "宋体", "Times New Roman", Pt(12))
        style.paragraph_format.line_spacing = Pt(20)

    if "EndNote Bibliography" in [style.name for style in doc.styles]:
        style = doc.styles["EndNote Bibliography"]
        set_style_fonts(style, "宋体", "Times New Roman", Pt(10.5))
        style.paragraph_format.line_spacing = Pt(20)

    if "toc 1" in [style.name for style in doc.styles]:
        style = doc.styles["toc 1"]
        set_style_fonts(style, "宋体", "Times New Roman", Pt(14))
        style.paragraph_format.first_line_indent = Pt(0)

    if "toc 2" in [style.name for style in doc.styles]:
        style = doc.styles["toc 2"]
        set_style_fonts(style, "宋体", "Times New Roman", Pt(10.5))
        style.paragraph_format.first_line_indent = Pt(0)


def apply_paragraph_formatting(doc: Document) -> dict[str, int]:
    counts = {
        "heading1": 0,
        "heading2": 0,
        "heading3": 0,
        "body": 0,
        "keywords": 0,
        "center_titles": 0,
        "declarations": 0,
        "captions": 0,
        "references": 0,
        "toc": 0,
        "list_items": 0,
    }

    list_mode = False
    list_headings = {"插图清单", "附表清单"}

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().replace("\xa0", " ")
        style_name = paragraph.style.name if paragraph.style else ""

        if text in list_headings:
            list_mode = True
        elif text in {"目   录", "目 录", "第一章 绪论"}:
            list_mode = False

        if text == "学位论文":
            apply_center_heading(paragraph, 18)
            counts["center_titles"] += 1
            continue
        if idx == 8 and text:
            apply_cover_title_paragraph(paragraph, cn=True)
            counts["center_titles"] += 1
            continue
        if idx == 9 and text:
            apply_cover_title_paragraph(paragraph, cn=False)
            counts["center_titles"] += 1
            continue

        if text in DECLARATION_HEADINGS:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                normalize_run_family_keep_size(run, "黑体", "Times New Roman")
            counts["declarations"] += 1
            continue

        if text in SPECIAL_CENTER_HEADINGS:
            apply_center_heading(paragraph, 16)
            counts["center_titles"] += 1
            continue

        if text.startswith(KEYWORD_PREFIXES):
            apply_keyword_paragraph(paragraph)
            counts["keywords"] += 1
            continue

        if style_name == "Heading 1":
            apply_heading_style(paragraph, 16, WD_ALIGN_PARAGRAPH.CENTER)
            counts["heading1"] += 1
            continue
        if style_name == "Heading 2":
            apply_heading_style(paragraph, 14, WD_ALIGN_PARAGRAPH.LEFT)
            counts["heading2"] += 1
            continue
        if style_name == "Heading 3":
            apply_heading_style(paragraph, 12, WD_ALIGN_PARAGRAPH.LEFT)
            counts["heading3"] += 1
            continue

        if style_name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_fonts(run, "黑体", "Times New Roman", Pt(10.5), bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)
            counts["captions"] += 1
            continue

        if style_name == "EndNote Bibliography":
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_fonts(run, "宋体", "Times New Roman", Pt(10.5))
            counts["references"] += 1
            continue

        if style_name in {"toc 1", "toc 2"}:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_run_fonts(
                    run,
                    "宋体",
                    "Times New Roman",
                    Pt(14 if style_name == "toc 1" else 10.5),
                )
            counts["toc"] += 1
            continue

        if list_mode and text:
            apply_list_item_paragraph(paragraph)
            counts["list_items"] += 1
            continue

        if idx <= 26 and text:
            apply_cover_meta_paragraph(paragraph)
            continue

        if looks_like_body_text(text):
            apply_body_paragraph(paragraph)
            counts["body"] += 1
            continue

        if style_name in {"Normal", "Body Text", "Body Text Indent"} and text:
            for run in paragraph.runs:
                normalize_run_family_keep_size(run, "宋体", "Times New Roman")

    return counts


def main() -> None:
    args = parse_args()
    src = Path(args.input).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    print(f"Input : {src}")
    print(f"Output: {output}")
    print("Step 1/5: repairing broken relationships if needed...")
    fixed_path = fix_docx_null_relationships(src)

    try:
        print("Step 2/5: loading document...")
        doc = Document(fixed_path)
        header_left_text = args.header_left or infer_header_left(doc)

        print("Step 3/5: applying CAU page setup and style definitions...")
        configure_page_setup(doc)
        configure_styles(doc)
        enable_update_fields_on_open(doc)

        print("Step 4/5: formatting paragraphs, tables, headers, and footers...")
        para_counts = apply_paragraph_formatting(doc)
        table_count = apply_table_formatting(doc)
        front_sections, main_sections = configure_headers_and_footers(doc, header_left_text)

        print("Step 5/5: saving...")
        doc.save(output)

        print("\nFormatting complete.")
        print(f"Header left text : {header_left_text}")
        print(f"Front sections   : {front_sections}")
        print(f"Main sections    : {main_sections}")
        print(f"Tables processed : {table_count}")
        for key, value in para_counts.items():
            print(f"{key:15}: {value}")
        print("\nWord will refresh page numbers and the directory field when the file is opened.")
    finally:
        try:
            os.remove(fixed_path)
        except OSError:
            pass


if __name__ == "__main__":
    main()
